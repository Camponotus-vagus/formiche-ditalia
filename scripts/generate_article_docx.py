#!/usr/bin/env python3
"""
Generate a Word document (.docx) for the Formiche d'Italia article
targeting Myrmecological News formatting guidelines.

Requirements: python-docx (pip install python-docx)
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Taxonomic names that must be italicized ──────────────────────────
ITALIC_TERMS = [
    # Order / Family
    "Hymenoptera", "Formicidae",
    # Subfamilies
    "Myrmicinae", "Formicinae", "Dolichoderinae", "Ponerinae",
    "Leptanillinae", "Proceratiinae", "Amblyoponinae",
    # Genera (all mentioned in the text)
    "Stigmatomma", "Proceratium", "Leptanilla", "Camponotus",
    "Mesquite",  # NOT a taxon, but the software — will exclude below
]

# Software names that should NOT be italicised even if they look like taxa
NOT_ITALIC = {"Mesquite"}

# Build regex: match whole words, longest first to avoid partial matches
ITALIC_TAXA = sorted(
    [t for t in ITALIC_TERMS if t not in NOT_ITALIC],
    key=len, reverse=True
)
TAXA_PATTERN = re.compile(r'\b(' + '|'.join(ITALIC_TAXA) + r')\b')

# Author surnames that should appear in small caps (bold fallback)
AUTHOR_SURNAMES_IN_CITATIONS = [
    "Baroni Urbani", "Hoelldobler", "Wilson", "Lebas", "Galkowski",
    "Blatrix", "Wegnez", "Maddison", "Mensa", "Norton", "Patterson",
    "Schneider", "Oeveraas", "Johansen", "Priem", "Piwowar", "Orr",
    "Schifani", "Vignes Lebbe", "Chesselet", "Diep Thi", "Zapparoli",
    "Dallwitz", "Paine", "Zurcher", "Rigato", "Mei",
]

OUTPUT_PATH = "/Users/francesco.mensa/Downloads/Progetto Formiche d'Italia/Mensa_FormicheDItalia_MyrmecolNews_DRAFT.docx"


def set_default_font(doc):
    """Set document-wide default font."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    # Set double spacing
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def configure_heading_styles(doc):
    """Configure heading styles to match journal requirements."""
    for level, size, bold in [(1, 14, True), (2, 12, True), (3, 12, True)]:
        style_name = f'Heading {level}'
        style = doc.styles[style_name]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(size)
        font.bold = bold
        font.color.rgb = RGBColor(0, 0, 0)
        if level >= 2:
            font.italic = False
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        pf.keep_with_next = True


def add_page_numbers(doc):
    """Add page numbers to the footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Page number field
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)


def add_formatted_paragraph(doc, text, style='Normal', alignment=None,
                            bold=False, italic=False, font_size=None,
                            space_after=None, space_before=None):
    """Add a paragraph with basic formatting, auto-italicising taxa."""
    p = doc.add_paragraph(style=style)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)

    # Split text by taxon names to italicise them
    parts = TAXA_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        run = p.add_run(part)
        run.font.name = 'Times New Roman'
        if font_size:
            run.font.size = Pt(font_size)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        # Auto-italicise taxonomic names
        if part in ITALIC_TAXA:
            run.italic = True
    return p


def add_mixed_paragraph(doc, segments, style='Normal', alignment=None,
                        space_after=None, space_before=None):
    """
    Add a paragraph from a list of (text, props) tuples.
    props is a dict with optional keys: bold, italic, font_size, superscript.
    Taxon names within each segment are auto-italicised.
    """
    p = doc.add_paragraph(style=style)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)

    for text, props in segments:
        props = props or {}
        # Split by taxa
        parts = TAXA_PATTERN.split(text)
        for part in parts:
            if not part:
                continue
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            if props.get('font_size'):
                run.font.size = Pt(props['font_size'])
            if props.get('bold'):
                run.bold = True
            if props.get('italic') or part in ITALIC_TAXA:
                run.italic = True
            if props.get('superscript'):
                run.font.superscript = True
    return p


def add_reference(doc, text):
    """Add a reference entry with hanging indent and auto-italicised taxa."""
    p = doc.add_paragraph(style='Normal')
    # Hanging indent: first line 0, left indent 0.5 inch
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)
    pf.space_after = Pt(0)

    # Split text by taxa
    parts = TAXA_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        run = p.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if part in ITALIC_TAXA:
            run.italic = True
    return p


def build_document():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    set_default_font(doc)
    configure_heading_styles(doc)
    add_page_numbers(doc)

    # ═══════════════════════════════════════════════════════════════════
    # FRONT MATTER
    # ═══════════════════════════════════════════════════════════════════

    # Title
    add_formatted_paragraph(
        doc,
        "Formiche d'Italia: a web-based interactive identification key for Italian ant genera and species of Rome with entropy-weighted multi-access scoring",
        style='Normal',
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        font_size=14,
        space_after=12,
    )

    # Author
    add_formatted_paragraph(
        doc,
        "Francesco Simone Mensa",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        font_size=12,
        space_after=6,
    )

    # Contact
    add_formatted_paragraph(
        doc,
        "Contact: Francesco Simone Mensa, formicheditalia@gmail.com",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        font_size=11,
        italic=True,
        space_after=12,
    )

    # Keywords
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(12)
    run_label = p.add_run("Keywords: ")
    run_label.bold = True
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(12)
    kw_text = "Formicidae, interactive identification key, multi-access key, error tolerance, entropy scoring, Italian ants, Rome, web platform, biodiversity informatics"
    parts = TAXA_PATTERN.split(kw_text)
    for part in parts:
        if not part:
            continue
        run = p.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if part in ITALIC_TAXA:
            run.italic = True

    # Page break before abstract
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # ABSTRACT
    # ═══════════════════════════════════════════════════════════════════

    doc.add_heading('Abstract', level=1)

    abstract = (
        "We present Formiche d'Italia (https://formicheditalia.it), a freely accessible, "
        "bilingual (Italian/English) web platform providing the first interactive multi-access "
        "identification key for all 39 ant genera present in the Italian peninsula and 75 species "
        "documented in the city of Rome. The platform is based on 40 morphological characters "
        "derived from the FormiKey thesis (Mensa 2017), whose character matrices were validated "
        "by F. Rigato in 2015. The identification algorithm introduces a novel entropy-weighted "
        "scoring system: each character selection is weighted by its Shannon information gain at "
        "the time of selection, and a subfamily-aware penalisation mechanism automatically "
        "down-ranks genera from subfamilies not matching the user's selections. Additional "
        "features include error tolerance (adjustable mismatch threshold), a best-character "
        "suggestion based on information gain, interactive impact prediction badges, a "
        "morphological glossary with 25 terms, and offline functionality via Progressive Web App "
        "technology. The platform also hosts an expert directory of 10 Italian myrmecologists "
        "with OpenAlex-enriched academic profiles. All specimen photographs (342+) are sourced "
        "from AntWeb (CC BY-SA 3.0), achieving 100% coverage for both genera and species. The "
        "platform was developed using AI-assisted software engineering (Claude, Anthropic) and "
        "is deployed as a static site on Vercel. We describe the technical architecture, the "
        "scoring algorithm, and discuss the advantages of the entropy-weighted approach compared "
        "to existing multi-access key tools such as Lucid, Xper3, and DELTA."
    )
    add_formatted_paragraph(doc, abstract)

    # ═══════════════════════════════════════════════════════════════════
    # INTRODUCTION
    # ═══════════════════════════════════════════════════════════════════

    doc.add_heading('Introduction', level=1)

    intro_p1 = (
        "Ants (Hymenoptera: Formicidae) are among the most ecologically important insect groups, "
        "functioning as primary predators of other invertebrates, soil engineers, and seed "
        "dispersers across virtually all terrestrial habitats (Hoelldobler & Wilson 1990). Italy "
        "hosts seven subfamilies (Myrmicinae, Formicinae, Dolichoderinae, Ponerinae, "
        "Leptanillinae, Proceratiinae, Amblyoponinae), encompassing 41 genera and over 267 "
        "species (Schifani 2022). The city of Rome alone harbours at least 75 species, "
        "representing approximately 30% of the peninsular fauna (Zapparoli 1997, Mei unpubl. "
        "data, Mensa 2017)."
    )
    add_formatted_paragraph(doc, intro_p1)

    intro_p2 = (
        "Despite this diversity, no modern, freely accessible digital identification tool "
        "existed for Italian Formicidae prior to the present work. Identification of ant genera "
        "and species has traditionally relied on dichotomous keys (e.g., Baroni Urbani 1971, "
        "Lebas & al. 2016), which impose a fixed sequence of character evaluations and are "
        "intolerant of user errors or missing observations. Multi-access (polyclave) keys "
        "address these limitations by allowing characters to be evaluated in any order, but "
        "existing software platforms \u2014 including Lucid (Norton & al. 2012), Xper3 (Vignes "
        "Lebbe & al. 2016), DELTA (Dallwitz & al. 1999), and Clavis (Oeveraas & Johansen "
        "2022) \u2014 present several drawbacks for the target use case: (1) they typically require "
        "installation or specialised software, (2) they lack built-in error tolerance "
        "mechanisms, (3) they do not weight characters by their discriminative power, and (4) "
        "they offer limited or no offline capability for field use."
    )
    add_formatted_paragraph(doc, intro_p2)

    intro_p3 = (
        "Formiche d'Italia was developed to address these gaps by providing a web-based, "
        "mobile-friendly identification platform with three principal innovations: (1) an "
        "entropy-weighted scoring algorithm that assigns higher importance to characters with "
        "greater discriminative power at the time of selection; (2) a subfamily-aware "
        "penalisation system that leverages the hierarchical structure of the character matrices "
        "to improve scoring accuracy; and (3) a comprehensive suite of user-guidance features "
        "including best-character suggestion, impact prediction, and a morphological glossary."
    )
    add_formatted_paragraph(doc, intro_p3)

    intro_p4 = (
        "The platform builds upon the FormiKey thesis (Mensa 2017), a master's thesis that "
        "produced NEXUS character matrices for all Italian ant genera and the species of Rome, "
        "with matrices validated by F. Rigato (Museo Civico di Storia Naturale di Milano) in "
        "2015."
    )
    add_formatted_paragraph(doc, intro_p4)

    # ═══════════════════════════════════════════════════════════════════
    # MATERIAL AND METHODS
    # ═══════════════════════════════════════════════════════════════════

    doc.add_heading('Material and methods', level=1)

    # -- Data sources --
    doc.add_heading('Data sources', level=2)

    meth_ds = (
        'The taxonomic backbone derives from the FormiKey thesis '
        '("\u0043hiave interattiva per il riconoscimento dei generi di Formicidae italiani '
        'e delle specie della citt\u00e0 di Roma"; Mensa 2017), produced at Roma Tre University '
        'under the supervision of Prof. A. Di Giulio. The thesis generated NEXUS-format '
        'character matrices using the software Mesquite (Maddison & Maddison 2023) covering: '
        '(a) a subfamily-level matrix (9 characters \u00d7 7 subfamilies); (b) four genus-level '
        'matrices (14\u201350 characters per subfamily); (c) species-level matrices for 18 genera '
        'present in Rome. Three additional genera (Stigmatomma, Proceratium, Leptanilla) are '
        'monotypic in Italy and were included without genus-level matrices. All genus-level '
        'matrices were validated by F. Rigato (MSNM) in 2015. The species checklist follows '
        'Schifani (2022). Specimen photographs were sourced from AntWeb (CC BY-SA 3.0).'
    )
    add_formatted_paragraph(doc, meth_ds)

    # -- Platform architecture --
    doc.add_heading('Platform architecture', level=2)

    meth_arch = (
        "Formiche d'Italia is implemented as a static website using the Astro framework (v5.x) "
        "with React components for client-side interactivity. The site generates 132 static HTML "
        "pages at build time, with seven React \"islands\" providing dynamic functionality. A "
        "Python data pipeline parses NEXUS matrices into JSON data files consumed at build time. "
        "The platform supports bilingual content (Italian/English) with 181 translation keys and "
        "offline functionality through a service worker (Progressive Web App). The platform was "
        "developed using AI-assisted software engineering (Claude, Anthropic), with all taxonomic "
        "data and scientific decisions made by the first author."
    )
    add_formatted_paragraph(doc, meth_arch)

    # -- Identification algorithm --
    doc.add_heading('Identification algorithm', level=2)

    meth_algo_intro = (
        "The identification algorithm introduces three innovations:"
    )
    add_formatted_paragraph(doc, meth_algo_intro)

    # Innovation 1
    algo1 = (
        "(1) Entropy-weighted character scoring: When a user selects a character state, "
        "the Shannon entropy H of that character across current candidate genera is calculated: "
        "H(c) = \u2013\u03a3i[p(si) \u00b7 log2(p(si))], where p(si) is the proportion of candidates "
        "exhibiting state si. This is normalised and stored as the weight wj. The weighted "
        "score for genus g is: score(g) = \u03a3j[match(g,j) \u00b7 wj] / \u03a3j[wj]."
    )
    add_formatted_paragraph(doc, algo1)

    # Innovation 2
    algo2 = (
        "(2) Subfamily-aware penalisation: When all selections share the same subfamily "
        "scope, out-of-scope genera with missing data receive a 0.8\u00d7 penalty (vs 0.3\u00d7 for "
        "in-scope), and genera with no data at all score 0.2."
    )
    add_formatted_paragraph(doc, algo2)

    # Innovation 3
    algo3 = (
        "(3) Best-character suggestion: The unused character with highest Shannon entropy "
        "among remaining candidates is highlighted, with impact prediction badges showing "
        "elimination counts per state."
    )
    add_formatted_paragraph(doc, algo3)

    # ═══════════════════════════════════════════════════════════════════
    # RESULTS
    # ═══════════════════════════════════════════════════════════════════

    doc.add_heading('Results', level=1)

    results = (
        "The platform (https://formicheditalia.it) provides: (a) an interactive identification "
        "key covering 39 Italian ant genera and 75 species of Rome, with 40 morphological "
        "characters classified by body region and difficulty; (b) individual pages for each "
        "genus and species with specimen photographs; (c) an expert directory of 10 Italian "
        "myrmecologists with OpenAlex-enriched profiles; (d) three editorial pages for "
        "non-specialist users; (e) a morphological glossary of 25 terms; (f) bilingual support "
        "and offline functionality. The platform achieves 100% photograph coverage."
    )
    add_formatted_paragraph(doc, results)

    # ═══════════════════════════════════════════════════════════════════
    # DISCUSSION
    # ═══════════════════════════════════════════════════════════════════

    doc.add_heading('Discussion', level=1)

    disc_p1 = (
        "The entropy-weighted scoring approach offers several advantages over binary elimination "
        "used by most existing tools. In traditional multi-access keys, selecting a state "
        "immediately eliminates non-matching genera, making the system fragile to errors. Our "
        "approach ranks genera by compatibility, with error tolerance retaining potentially "
        "matching genera."
    )
    add_formatted_paragraph(doc, disc_p1)

    disc_p2 = (
        "The subfamily-aware penalisation is, to our knowledge, novel. It leverages hierarchical "
        "matrix structure without requiring users to navigate a taxonomic hierarchy."
    )
    add_formatted_paragraph(doc, disc_p2)

    disc_p3 = (
        "Current limitations include species coverage limited to Rome (75 species) and matrices "
        "validated in 2015. Expansion to all 267+ Italian species is planned. Future features "
        "include a simplified mode for beginners, side-by-side comparison, AI photo "
        "pre-filtering, and GBIF integration."
    )
    add_formatted_paragraph(doc, disc_p3)

    # ═══════════════════════════════════════════════════════════════════
    # ACKNOWLEDGEMENTS
    # ═══════════════════════════════════════════════════════════════════

    doc.add_heading('Acknowledgements', level=1)

    ack = (
        "We thank F. Rigato (Museo Civico di Storia Naturale di Milano) for validation of the "
        "morphological character matrices. Specimen photographs are from AntWeb (CC BY-SA 3.0). "
        "Expert data from OpenAlex API. Species checklist follows Schifani (2022). Web platform "
        "development assisted by Claude (Anthropic)."
    )
    add_formatted_paragraph(doc, ack)

    # ═══════════════════════════════════════════════════════════════════
    # REFERENCES
    # ═══════════════════════════════════════════════════════════════════

    doc.add_heading('References', level=1)

    references = [
        "Baroni Urbani, C. 1971: Catalogo delle specie di Formicidae d'Italia. \u2013 Memorie della Societ\u00e0 Entomologica Italiana 50: 5-287.",
        "Dallwitz, M.J., Paine, T.A. & Zurcher, E.J. 1999: User's guide to the DELTA Editor. \u2013 <https://www.delta-intkey.com/>, retrieved on 03 April 2026.",
        "Hoelldobler, B. & Wilson, E.O. 1990: The ants. \u2013 Belknap Press of Harvard University Press, Cambridge, MA, 732 pp.",
        "Lebas, C., Galkowski, C., Blatrix, R. & Wegnez, P. 2016: Fourmis d'Europe occidentale. \u2013 Delachaux et Niestl\u00e9, Paris, 415 pp.",
        "Maddison, W.P. & Maddison, D.R. 2023: Mesquite: a modular system for evolutionary analysis, version 3.81. \u2013 <http://www.mesquiteproject.org>, retrieved on 03 April 2026.",
        "Mensa, F.S. 2017: Chiave interattiva per il riconoscimento dei generi di Formicidae italiani e delle specie della citt\u00e0 di Roma. \u2013 Master's thesis, Roma Tre University, Rome, 69 pp.",
        "Norton, G.A., Patterson, D.J. & Schneider, M. 2012: LucidMobile: An application for interactive identification of organisms. \u2013 Biodiversity Informatics 8: 43-47.",
        "Oeveraas, H. & Johansen, V. 2022: Clavis \u2013 A web-based identification key format. \u2013 Biodiversity Data Journal 10: e80517.",
        "Priem, J., Piwowar, H. & Orr, R. 2022: OpenAlex: A fully-open index of scholarly works, authors, venues, institutions, and concepts. \u2013 ArXiv preprint arXiv:2205.01833.",
        "Schifani, E. 2022: Checklist of the Italian Fauna \u2013 Formicidae. \u2013 [Reference to be completed].",
        "Vignes Lebbe, R., Chesselet, P. & Diep Thi, M.H. 2016: Xper3: new tools for collaborating, training and transmitting knowledge. \u2013 Botany Letters 163: 463-467.",
        "Zapparoli, M. 1997: Urban development and insect biodiversity of the Rome area, Italy. \u2013 Landscape and Urban Planning 38: 77-86.",
    ]

    for ref in references:
        add_reference(doc, ref)

    # ═══════════════════════════════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════════════════════════════

    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == '__main__':
    build_document()
